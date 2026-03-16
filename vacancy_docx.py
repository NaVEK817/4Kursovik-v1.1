from __future__ import annotations

from datetime import datetime
from pathlib import Path

try:
    from docx import Document
except Exception:  # noqa: BLE001
    Document = None


def save_vacancy_to_docx(text: str, filename: Path | None = None) -> Path:
    """Сохранить текст вакансии в оформленный DOCX‑файл.

    По умолчанию создаёт файл вида vacancy_YYYYMMDD_HHMMSS.docx в директории проекта.
    """
    if Document is None:
        raise RuntimeError(
            "Библиотека python-docx не установлена. "
            "Установите её командой: python3 -m pip install python-docx"
        )

    base_dir = Path(__file__).resolve().parent
    if filename is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = base_dir / f"vacancy_{ts}.docx"

    doc = Document()

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        lines = ["Вакансия"]

    # Первая непустая строка как заголовок
    title = lines[0]
    body_lines = lines[1:] or []

    doc.add_heading(title, level=1)

    for line in body_lines:
        doc.add_paragraph(line)

    doc.save(filename)
    return filename

