# -*- coding: utf-8 -*-
"""
Окно создания оффера для кандидата с автоматическим назначением собеседования
"""
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                             QLineEdit, QPushButton, QTextEdit, 
                             QMessageBox, QFileDialog, QGroupBox, QDateEdit,
                             QSpinBox, QFormLayout, QCheckBox, QComboBox)
from PyQt5.QtCore import Qt, QDate
import styles

# Импортируем планировщик
from ai_offer_scheduler import AIOfferScheduler


class DocumentWindow(QWidget):
    """Окно создания оффера для кандидата с автоназначением собеседования"""
    
    def __init__(self, vacancy, candidate):
        super().__init__()
        self.vacancy = vacancy
        self.candidate = candidate
        self.scheduler = AIOfferScheduler()
        self.interview_scheduled = False
        self.interview_data = None
        self.init_ui()
        
    def init_ui(self):
        """Инициализация интерфейса"""
        self.setWindowTitle("S7 Recruitment - Создание оффера и назначение собеседования")
        self.setGeometry(200, 200, 650, 800)
        self.setStyleSheet(styles.MAIN_STYLE)
        
        # Основной layout
        layout = QVBoxLayout()
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Заголовок
        header_label = QLabel("📄 Создание оффера и назначение собеседования")
        header_label.setObjectName("headerLabel")
        layout.addWidget(header_label)
        
        # Информация о кандидате
        candidate_group = QGroupBox("Информация о кандидате")
        candidate_layout = QFormLayout()
        
        # Формируем ФИО
        full_name = self.get_candidate_full_name()
        
        name_label = QLabel(full_name)
        name_label.setStyleSheet(f"font-weight: bold; color: {styles.S7_GREEN};")
        candidate_layout.addRow("👤 ФИО:", name_label)
        
        # Желаемая должность
        desired_title = self.candidate.get('title', 'Не указана')
        candidate_layout.addRow("💼 Желаемая должность:", QLabel(desired_title))
        
        # Город
        city = self.candidate.get('area', 'Не указан')
        candidate_layout.addRow("🏙️ Город:", QLabel(city))
        
        # Контакты
        phone = self.candidate.get('phone', 'Не указан')
        email = self.candidate.get('email', 'Не указан')
        candidate_layout.addRow("📞 Телефон:", QLabel(phone))
        candidate_layout.addRow("✉️ Email:", QLabel(email))
        
        candidate_group.setLayout(candidate_layout)
        layout.addWidget(candidate_group)
        
        # Информация о вакансии
        vacancy_group = QGroupBox("Информация о вакансии")
        vacancy_layout = QFormLayout()
        
        vacancy_title = QLabel(self.vacancy.get('title', 'Не указана'))
        vacancy_title.setStyleSheet(f"font-weight: bold; color: {styles.S7_GREEN};")
        vacancy_layout.addRow("📋 Вакансия:", vacancy_title)
        
        vacancy_city = self.vacancy.get('area', 'Не указан')
        vacancy_layout.addRow("🏙️ Город работы:", QLabel(vacancy_city))
        
        vacancy_salary = self.vacancy.get('salary', 'Не указана')
        vacancy_layout.addRow("💰 Зарплата:", QLabel(vacancy_salary))
        
        vacancy_schedule = self.vacancy.get('schedule', 'Не указан')
        vacancy_layout.addRow("⏰ График:", QLabel(vacancy_schedule))
        
        vacancy_group.setLayout(vacancy_layout)
        layout.addWidget(vacancy_group)
        
        # Параметры оффера
        offer_group = QGroupBox("Параметры оффера")
        offer_layout = QFormLayout()
        
        # Дата начала работы
        self.start_date = QDateEdit()
        self.start_date.setDate(QDate.currentDate().addDays(14))
        self.start_date.setCalendarPopup(True)
        offer_layout.addRow("📅 Предполагаемая дата выхода:", self.start_date)
        
        # Испытательный срок
        self.probation_period = QSpinBox()
        self.probation_period.setRange(1, 6)
        self.probation_period.setValue(3)
        self.probation_period.setSuffix(" месяца")
        offer_layout.addRow("⏱️ Испытательный срок:", self.probation_period)
        
        # Зарплатное предложение
        self.salary_offer = QLineEdit()
        self.salary_offer.setPlaceholderText("Например: 80 000 руб.")
        vacancy_salary = self.vacancy.get('salary', '')
        if vacancy_salary:
            self.salary_offer.setText(vacancy_salary)
        offer_layout.addRow("💰 Предложение по зарплате:", self.salary_offer)
        
        offer_group.setLayout(offer_layout)
        layout.addWidget(offer_group)
        
        # БЛОК НАЗНАЧЕНИЯ СОБЕСЕДОВАНИЯ
        interview_group = QGroupBox("📅 Назначение собеседования")
        interview_layout = QVBoxLayout()
        
        # Чекбокс для автоматического назначения
        self.auto_schedule_check = QCheckBox("✅ Автоматически назначить собеседование после создания оффера")
        self.auto_schedule_check.setChecked(True)
        self.auto_schedule_check.setStyleSheet(f"color: {styles.S7_GREEN}; font-weight: bold;")
        interview_layout.addWidget(self.auto_schedule_check)
        
        # Параметры собеседования
        interview_params_layout = QFormLayout()
        
        # Интервьюер
        self.interviewer_input = QLineEdit()
        self.interviewer_input.setPlaceholderText("Кто проводит собеседование")
        # Пытаемся получить имя пользователя из системы
        import getpass
        import os
        recruiter_name = os.environ.get('USERNAME', getpass.getuser())
        self.interviewer_input.setText(recruiter_name)
        interview_params_layout.addRow("👥 Интервьюер:", self.interviewer_input)
        
        # Предпочтительное время
        time_layout = QHBoxLayout()
        self.preferred_time_combo = QComboBox()
        self.preferred_time_combo.addItems(["Утро (9-12)", "День (12-15)", "Вечер (15-18)", "Любое"])
        time_layout.addWidget(self.preferred_time_combo)
        interview_params_layout.addRow("⏰ Предпочтительное время:", time_layout)
        
        interview_layout.addLayout(interview_params_layout)
        
        # Кнопка для ручного назначения
        self.manual_schedule_btn = QPushButton("📅 Назначить собеседование вручную")
        self.manual_schedule_btn.clicked.connect(self.manual_schedule_interview)
        self.manual_schedule_btn.setCursor(Qt.PointingHandCursor)
        interview_layout.addWidget(self.manual_schedule_btn)
        
        # Информация о назначенном собеседовании
        self.interview_info_label = QLabel()
        self.interview_info_label.setWordWrap(True)
        self.interview_info_label.setVisible(False)
        interview_layout.addWidget(self.interview_info_label)
        
        interview_group.setLayout(interview_layout)
        layout.addWidget(interview_group)
        
        # Дополнительные условия
        conditions_group = QGroupBox("Дополнительные условия")
        conditions_layout = QVBoxLayout()
        
        self.conditions_text = QTextEdit()
        self.conditions_text.setMaximumHeight(80)
        self.conditions_text.setPlaceholderText("ДМС, бонусы, дополнительные льготы...")
        conditions_layout.addWidget(self.conditions_text)
        
        conditions_group.setLayout(conditions_layout)
        layout.addWidget(conditions_group)
        
        # Комментарий
        comment_group = QGroupBox("Комментарий для кандидата")
        comment_layout = QVBoxLayout()
        
        self.comment_text = QTextEdit()
        self.comment_text.setMaximumHeight(60)
        self.comment_text.setPlaceholderText("Персональное обращение к кандидату...")
        comment_layout.addWidget(self.comment_text)
        
        comment_group.setLayout(comment_layout)
        layout.addWidget(comment_group)
        
        # Прикрепление резюме
        resume_group = QGroupBox("📎 Прикрепление резюме")
        resume_layout = QHBoxLayout()
        
        self.resume_path_label = QLabel("Файл не выбран")
        self.resume_path_label.setStyleSheet(f"color: {styles.S7_GRAY};")
        resume_layout.addWidget(self.resume_path_label, 1)
        
        self.attach_resume_btn = QPushButton("📎 Выбрать файл")
        self.attach_resume_btn.clicked.connect(self.attach_resume_file)
        self.attach_resume_btn.setCursor(Qt.PointingHandCursor)
        resume_layout.addWidget(self.attach_resume_btn)
        
        resume_group.setLayout(resume_layout)
        layout.addWidget(resume_group)
        
        # Кнопки
        buttons_layout = QHBoxLayout()
        
        self.create_btn = QPushButton("📄 Создать оффер и назначить собеседование")
        self.create_btn.clicked.connect(self.create_offer_and_schedule)
        self.create_btn.setCursor(Qt.PointingHandCursor)
        self.create_btn.setMinimumHeight(40)
        buttons_layout.addWidget(self.create_btn)
        
        buttons_layout.addStretch()
        
        self.cancel_btn = QPushButton("✖ Отмена")
        self.cancel_btn.clicked.connect(self.close)
        self.cancel_btn.setCursor(Qt.PointingHandCursor)
        buttons_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(buttons_layout)
        
        self.setLayout(layout)
    
    def get_candidate_full_name(self):
        """Возвращает полное имя кандидата"""
        if 'first_name' in self.candidate and 'last_name' in self.candidate:
            return f"{self.candidate.get('last_name', '')} {self.candidate.get('first_name', '')} {self.candidate.get('middle_name', '')}".strip()
        return self.candidate.get('name', 'Кандидат')
    
    def attach_resume_file(self):
        """Прикрепляет файл с резюме"""
        filename, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите файл с резюме",
            "",
            "PDF Files (*.pdf);;Word Documents (*.docx);;Text Files (*.txt);;All files (*.*)"
        )
        
        if filename:
            self.resume_path_label.setText(f"✅ {os.path.basename(filename)}")
            self.resume_path_label.setStyleSheet(f"color: {styles.S7_GREEN};")
            self.resume_file = filename
    
    def manual_schedule_interview(self):
        """Ручное назначение собеседования"""
        try:
            candidate_name = self.get_candidate_full_name()
            recruiter_name = self.interviewer_input.text().strip()
            
            if not recruiter_name:
                recruiter_name = "HR-менеджер"
            
            # Вызываем планировщик
            result = self.scheduler.schedule_interview_for_offer(
                candidate_data=self.candidate,
                vacancy_data=self.vacancy,
                recruiter_name=recruiter_name,
                days_ahead=14
            )
            
            if result.get('success'):
                self.interview_scheduled = True
                self.interview_data = result.get('interview', {})
                
                self.interview_info_label.setText(
                    f"✅ Собеседование назначено на {self.interview_data.get('date')} в {self.interview_data.get('time')}\n"
                    f"👥 Интервьюер: {recruiter_name}"
                )
                self.interview_info_label.setStyleSheet(f"color: {styles.S7_GREEN}; padding: 5px;")
                self.interview_info_label.setVisible(True)
                
                QMessageBox.information(
                    self,
                    "Собеседование назначено",
                    f"✅ Собеседование для кандидата {candidate_name} успешно назначено!\n\n"
                    f"📅 Дата: {self.interview_data.get('date')}\n"
                    f"⏰ Время: {self.interview_data.get('time')}\n"
                    f"👥 Интервьюер: {recruiter_name}"
                )
            else:
                QMessageBox.warning(self, "Ошибка", result.get('message', 'Не удалось назначить собеседование'))
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось назначить собеседование: {str(e)}")
    
    def create_offer_and_schedule(self):
        """Создание оффера и назначение собеседования"""
        # Проверка заполнения
        if not self.salary_offer.text().strip():
            QMessageBox.warning(self, "Предупреждение", "Укажите зарплатное предложение")
            return
        
        candidate_name = self.get_candidate_full_name()
        
        # Если автоназначение включено и еще не назначено
        if self.auto_schedule_check.isChecked() and not self.interview_scheduled:
            recruiter_name = self.interviewer_input.text().strip()
            if not recruiter_name:
                recruiter_name = "HR-менеджер"
            
            # Назначаем собеседование
            result = self.scheduler.schedule_interview_for_offer(
                candidate_data=self.candidate,
                vacancy_data=self.vacancy,
                recruiter_name=recruiter_name,
                days_ahead=14
            )
            
            if result.get('success'):
                self.interview_scheduled = True
                self.interview_data = result.get('interview', {})
            else:
                reply = QMessageBox.question(
                    self,
                    "Ошибка назначения",
                    f"Не удалось автоматически назначить собеседование: {result.get('message')}\n\n"
                    f"Продолжить создание оффера без назначения собеседования?",
                    QMessageBox.Yes | QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return
        
        # Создаем оффер
        self.create_offer(candidate_name)
    
    def create_offer(self, candidate_name):
        """Создание документа-оффера"""
        # Выбор места сохранения
        filename, _ = QFileDialog.getSaveFileName(
            self, 
            "Сохранить оффер",
            f"Оффер_{candidate_name}_{self.vacancy.get('title', '')[:30]}.docx",
            "Word Documents (*.docx)"
        )
        
        if not filename:
            return
        
        try:
            # Создание документа
            doc = Document()
            
            # Установка стилей
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Заголовок
            title = doc.add_heading('ПРЕДЛОЖЕНИЕ О РАБОТЕ (JOB OFFER)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(0, 166, 81)
            
            doc.add_paragraph()
            
            # Дата
            date_para = doc.add_paragraph(f'г. Москва, {datetime.now().strftime("%d.%m.%Y")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()
            
            # Обращение
            doc.add_paragraph(f'Уважаемый(ая) {candidate_name}!')
            doc.add_paragraph()
            
            # Основной текст
            doc.add_paragraph(
                'Мы рады предложить Вам присоединиться к команде S7 Airlines в качестве '
                f'**{self.vacancy.get("title", "сотрудника")}**.'
            )
            doc.add_paragraph()
            
            # Условия работы
            doc.add_heading('Условия работы:', level=1)
            
            conditions_table = doc.add_table(rows=5, cols=2)
            conditions_table.style = 'Light Grid Accent 1'
            
            # Заполняем таблицу
            cells = conditions_table.rows[0].cells
            cells[0].text = 'Должность'
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[1].text = self.vacancy.get('title', 'Не указана')
            
            cells = conditions_table.rows[1].cells
            cells[0].text = 'Место работы'
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[1].text = self.vacancy.get('area', 'Не указано')
            
            cells = conditions_table.rows[2].cells
            cells[0].text = 'График работы'
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[1].text = self.vacancy.get('schedule', 'Не указан')
            
            cells = conditions_table.rows[3].cells
            cells[0].text = 'Предполагаемая дата выхода'
            cells[0].paragraphs[0].runs[0].font.bold = True
            cells[1].text = self.start_date.date().toString("dd.MM.yyyy")
            
            cells = conditions_table.rows[4].cells
            cells[0].text = 'Собеседование'
            cells[0].paragraphs[0].runs[0].font.bold = True
            if self.interview_scheduled and self.interview_data:
                cells[1].text = f"{self.interview_data.get('date')} в {self.interview_data.get('time')}"
            else:
                cells[1].text = "Будет согласовано дополнительно"
            
            doc.add_paragraph()
            
            # Компенсация
            doc.add_heading('Компенсационный пакет:', level=1)
            
            salary_text = f'• Ежемесячная заработная плата: {self.salary_offer.text()} до вычета НДФЛ'
            doc.add_paragraph(salary_text, style='List Bullet')
            
            probation_text = f'• Испытательный срок: {self.probation_period.value()} месяца'
            doc.add_paragraph(probation_text, style='List Bullet')
            
            # Дополнительные условия
            conditions = self.conditions_text.toPlainText().strip()
            if conditions:
                doc.add_paragraph()
                doc.add_heading('Дополнительные льготы и преимущества:', level=1)
                for line in conditions.split('\n'):
                    if line.strip():
                        doc.add_paragraph(f'• {line.strip()}', style='List Bullet')
            
            # Стандартные условия S7
            doc.add_paragraph()
            doc.add_heading('Стандартные условия для сотрудников S7:', level=1)
            doc.add_paragraph('• ДМС после испытательного срока (диагностика, лечение, стоматология со скидкой до 75%)', style='List Bullet')
            doc.add_paragraph('• Корпоративные тарифы на авиабилеты для Вас и членов семьи', style='List Bullet')
            doc.add_paragraph('• Программа корпоративных скидок PrimeZone (скидки до 80% у 1500+ партнеров)', style='List Bullet')
            doc.add_paragraph('• Программа психологической, финансовой и юридической поддержки "Понимаю"', style='List Bullet')
            
            # Информация о прикрепленных файлах
            if hasattr(self, 'resume_file'):
                doc.add_paragraph()
                doc.add_heading('Прикрепленные файлы:', level=1)
                doc.add_paragraph(f'• Резюме кандидата: {os.path.basename(self.resume_file)}', style='List Bullet')
            
            # Комментарий
            comment = self.comment_text.toPlainText().strip()
            if comment:
                doc.add_paragraph()
                doc.add_heading('Дополнительно:', level=1)
                doc.add_paragraph(comment)
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Подпись
            doc.add_paragraph('С уважением,')
            doc.add_paragraph('Команда S7 Recruitment')
            doc.add_paragraph()
            doc.add_paragraph('______________________          ______________________')
            doc.add_paragraph('      (подпись)                              (дата)')
            
            # Сохранение документа
            doc.save(filename)
            
            # Формируем сообщение об успехе
            success_message = f"✅ Оффер успешно создан и сохранен:\n{filename}\n\n"
            
            if self.interview_scheduled and self.interview_data:
                success_message += (
                    f"📅 Собеседование назначено на {self.interview_data.get('date')} в {self.interview_data.get('time')}\n"
                    f"👥 Интервьюер: {self.interviewer_input.text().strip()}\n"
                )
            
            if hasattr(self, 'resume_file'):
                success_message += f"📎 Прикреплено резюме: {os.path.basename(self.resume_file)}\n"
            
            QMessageBox.information(self, "Успех", success_message)
            
            self.close()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось создать оффер: {str(e)}")